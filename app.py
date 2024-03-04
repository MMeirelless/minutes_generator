from flask import Flask, render_template, url_for, redirect, request, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, BooleanField, SubmitField, SelectField, ValidationError
from wtforms.validators import DataRequired, Length, Email, EqualTo
from werkzeug.security import generate_password_hash, check_password_hash
from generator import report_generator, dict_to_html
from io import BytesIO
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = '\x9a\xd0lTA^~w?\xeb\xb1\xba\xf0\x98\xacbs\xd3\xc2dD\x04sq'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///site.db'
db = SQLAlchemy(app)
doc_title = None
output = None

login_manager = LoginManager(app)
login_manager.login_view = 'login'

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)
    plan = db.Column(db.String(20), nullable=False, default='Grátis')

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

class RegistrationForm(FlaskForm):
    username = StringField('Nome', validators=[DataRequired(), Length(min=2, max=20)])
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Senha', validators=[DataRequired()])
    confirm_password = PasswordField('Confirmar Senha', validators=[DataRequired(), EqualTo('password')])
    plan = SelectField('Plano', choices=[('Grátis', 'Grátis'), ('Pro', 'Pro'), ('Empresarial', 'Empresarial')])
    submit = SubmitField('Cadastrar-se')
    
    def validate_email(self, email):
        user = User.query.filter_by(email=email.data).first()
        if user:
            raise ValidationError('Email already registered. Please use a different email address.')


class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Senha', validators=[DataRequired()])
    remember = BooleanField('Lembrar acesso')
    submit = SubmitField('Entrar')

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('new_report'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user and check_password_hash(user.password, form.password.data):
            login_user(user, remember=form.remember.data)
            flash(f'Olá, {current_user.username}!', 'success')
            return redirect(url_for('new_report'))

        else:
            flash('Não foi possível entrar. Por favor, verifique seu e-mail e senha.', 'danger')

    return render_template('login.html', title='Login', form=form)

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    form = RegistrationForm()
    if form.validate_on_submit():
        existing_user = User.query.filter_by(email=form.email.data).first()
        if existing_user:
            flash('Email já em uso. Por favor, escolha outro.', 'danger')
            return redirect(url_for('register'))
        
        hashed_password = generate_password_hash(form.password.data)
        user = User(username=form.username.data, email=form.email.data, password=hashed_password, plan=form.plan.data)
        db.session.add(user)
        db.session.commit()
        flash('Sua conta foi criada! Bem vindo!', 'success')
        return redirect(url_for('login'))
    return render_template('register.html', title='Register', form=form)


@app.route('/logout')
def logout():
    logout_user()
    flash('Você saiu, até logo :)', 'danger')
    return redirect(url_for('home'))

@app.route("/plans")
def plans():
    return render_template('plans.html')

@app.route("/support")
def support():
    return render_template('support.html')

@app.route("/my_account")
def my_account():
    return render_template('dashboard/my_account.html')

@app.route("/my_reports")
def my_reports():
    return render_template('dashboard/my_reports.html')

@app.route('/new_report', methods=['GET', 'POST'])
def new_report():
    global output
    global doc_title
    html_content = ""
    if request.method == 'POST':
        if 'audio_file' not in request.files:
            flash('Por favor, selecione um arquivo.', 'danger')
            return redirect(request.url)
        
        audio_file = request.files['audio_file']

        if audio_file.filename == '':
            flash('Por favor, selecione um arquivo.', 'danger')
            return redirect(request.url)
        
        if audio_file and audio_file.filename.endswith('.mp3'):
            doc = Document()
            result = report_generator(audio_file)  
            minutes = result[0]
            doc_title = f"{result[1]}.docx"
            download_url = url_for('report_download')
            html_content = dict_to_html(minutes, download_url)
            
            for key, value in minutes.items():
                heading = ' '.join(word.capitalize() for word in key.split('_'))
                doc.add_heading(heading, level=1)
                doc.add_paragraph(value)
                doc.add_paragraph()

            output = BytesIO()
            doc.save(output)
            output.seek(0)

            flash('Relatório gerado com sucesso!', 'success')
            return render_template('dashboard/new_report.html', html_content=html_content,is_report_generated="true")
        
    return render_template('dashboard/new_report.html', html_content=html_content,is_report_generated="false")

@app.route('/report_download')
def report_download():
    global output
    global doc_title
    try:
        return send_file(output, download_name=doc_title, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        flash('Erro ao baixar o arquivo.', 'danger')
        return redirect(url_for('new_report'))


@app.route("/trash")
def trash():
    return render_template('dashboard/trash.html')

if __name__ == '__main__':
    with app.app_context():
        db.create_all()

    app.run(debug=True, host="0.0.0.0", port="9099")

