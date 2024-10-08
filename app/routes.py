from io import BytesIO
from docx import Document
from flask import Blueprint, session, render_template, redirect, url_for, request, flash, send_file
from flask_login import login_user, logout_user, login_required, current_user
from . import db
from .models import User, Report, Trash
from .forms import LoginForm, RegistrationForm, TwoFactorForm
from .services import report_generator, dict_to_html, send_verification_email
from werkzeug.security import generate_password_hash, check_password_hash
from json import dumps 
from datetime import datetime, timedelta
from hashlib import md5
from sqlalchemy.exc import IntegrityError
from random import randint, choices
from string import digits
import pyotp

main = Blueprint('main', __name__)
update_account_verification_code = "0"
code_is_sent = False

@main.route('/')
def home():
    if current_user.is_authenticated:
        return redirect(url_for('main.new_report'))
    else:
        return render_template('home.html')

@main.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('main.new_report'))

    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user and check_password_hash(user.password, form.password.data):
            if user.two_factor_secret:
                return render_template('verify_2fa.html', user_id=user.id, form=TwoFactorForm())
            else:
                login_user(user, remember=form.remember.data)
                flash(f'Olá, {user.username}!', 'success')
                return redirect(url_for('main.new_report'))
        else:
            flash('Não foi possível entrar. Por favor, verifique seu e-mail e senha.', 'danger')

    return render_template('login.html', title='Login', form=form)

@main.route('/verify_2fa', methods=['GET', 'POST'])
def verify_2fa():
    form = TwoFactorForm()
    if form.validate_on_submit():
        user_id = request.form.get('user_id')
        user = User.query.get(user_id)
        print(pyotp.TOTP(user.two_factor_secret))
        if user and pyotp.TOTP(user.two_factor_secret).verify(form.token.data):
            login_user(user, remember=True)
            return redirect(url_for('main.new_report'))
        else:
            flash('Token inválido.', 'danger')
            return render_template('verify_2fa.html', user_id=user_id, form=TwoFactorForm())
        
    return render_template('verify_2fa.html', form=form)

@main.route('/setup_2fa')
def setup_2fa():
    if not current_user.is_authenticated:
        return redirect(url_for('main.login'))

    secret = pyotp.random_base32()
    current_user.two_factor_secret = secret
    db.session.commit()

    url = pyotp.totp.TOTP(secret).provisioning_uri(current_user.email, issuer_name='Keysession')
    return render_template('setup_2fa.html', secret=secret, url=url)

@main.route('/register', methods=['GET', 'POST'])
def register():
    form = RegistrationForm()
    if current_user.is_authenticated:
        return redirect(url_for('main.new_report'))
    
    if request.method == 'POST':
        if form.validate_on_submit():
            verification_code = ''.join(choices(digits, k=6))
            session['email'] = form.email.data
            session['verification_code'] = verification_code
            session['username'] = form.username.data
            session['password'] = generate_password_hash(form.password.data)
            session['plan'] = form.plan.data
            session['user_pic'] = form.user_pic.data if form.user_pic.data != "" else f"user{randint(1, 6)}.webp"

            send_verification_email(form.email.data, verification_code)

            return redirect(url_for('main.verify_email'))
    return render_template('register.html', form=form)

@main.route('/verify_email', methods=['GET', 'POST'])
def verify_email():
    if current_user.is_authenticated:
        return redirect(url_for('main.new_report'))
    
    if request.method == 'POST':
        user_code = request.form['verification_code']

        if user_code == session.get('verification_code'):
            try:
                existing_user = User.query.filter_by(email=session.get('email')).first()
                if existing_user:
                    flash('Email já em uso. Por favor, escolha outro.', 'danger')
                    return redirect(url_for('main.register'))
                
                user = User(username=session.get('username'), email=session.get('email'), password=session.get('password'), plan=session.get('plan'), user_pic=session.get('user_pic'))
                db.session.add(user)
                db.session.commit()

                session.pop('email', None)
                session.pop('username', None)
                session.pop('password', None)
                session.pop('plan', None)
                session.pop('user_pic', None)
                session.pop('verification_code', None)

                flash('Sua conta foi criada! Bem vindo!', 'success')
                return redirect(url_for('main.login'))
            
            except IntegrityError:
                db.session.rollback()
                flash('Erro ao criar conta, por favor, tente novamente.', 'danger')
                return redirect(url_for('main.register'))
                    
        else:
            flash('Código de verificação inválido.', 'danger')
            return redirect(url_for('main.verify_email'))
        
    return render_template('verify_email.html')

@main.route('/logout')
def logout():
    logout_user()
    flash('Você saiu, até logo :)', 'danger')
    return redirect(url_for('main.home'))

@main.route("/plans")
def plans():
    return render_template('plans.html')

@main.route("/my_account")
def my_account():
    if current_user.is_authenticated:
        global update_account_verification_code
        global code_is_sent

        update_account_verification_code = "0"
        code_is_sent = False
        return render_template('dashboard/my_account.html')
    else:
        return redirect(url_for('main.home'))

@main.route('/my_reports')
def my_reports():
    if current_user.is_authenticated:
        reports = Report.query.filter_by(user_id=current_user.id).all()
        return render_template('dashboard/my_reports.html', reports=reports)
    else:
        return redirect(url_for('main.home'))

@main.route('/new_report', methods=['GET', 'POST'])
def new_report():
    if current_user.is_authenticated:
        global output
        global doc_title
        html_content_title = "Novo Relatório"
        html_content = ""

        # Cleaning Trash
        user_trashes = Trash.query.filter_by(user_id=current_user.id).all()
        for trash in user_trashes:
            if trash.date < datetime.now()-timedelta(days=7):
                print(datetime.now()-trash.date)
                db.session.delete(trash)
                db.session.commit()
        
        if request.method == 'POST':
            if 'audio_file' not in request.files:
                flash('Por favor, selecione um arquivo.', 'danger')
                return redirect(request.url)
            
            audio_file = request.files['audio_file']

            if audio_file.filename == '':
                flash('Por favor, selecione um arquivo.', 'danger')
                return redirect(request.url)
            
            if audio_file and (audio_file.filename.endswith('.mp3') or audio_file.filename.endswith('.mp4')):
                doc = Document()
                result = report_generator(audio_file)  
                minutes = result[0]
                doc_title = f"{result[1]}.docx"
                download_url = url_for('main.report_download')
                html_content_title = doc_title.replace("_", " ").split(".")[0]
                html_content = dict_to_html(minutes, download_url)

                for key, value in minutes.items():
                    heading = ' '.join(word.capitalize() for word in key.split('_'))
                    doc.add_heading(heading, level=1)
                    doc.add_paragraph(value)
                    doc.add_paragraph()

                output = BytesIO()
                doc.save(output)
                output.seek(0)

                flash('Relatório gerado com sucesso!', 'success')
                return render_template('dashboard/new_report.html', html_content=html_content,is_report_generated="true",html_content_title=html_content_title)
            
        return render_template('dashboard/new_report.html', html_content=html_content,is_report_generated="false",html_content_title=html_content_title)
    else:
        return redirect(url_for('main.home'))

@main.route('/report_download')
def report_download():
    if current_user.is_authenticated:
        global output
        global doc_title
        try:
            return send_file(output, download_name=doc_title, as_attachment=True, mimetype='mainlication/vnd.openxmlformats-officedocument.wordprocessingml.document')
        except Exception as e:
            flash('Erro ao baixar o arquivo.', 'danger')
            return redirect(url_for('main.new_report'))
    
    else:
        return redirect(url_for('main.home'))

@main.route('/save_report', methods=["POST"])
def save_report():
    if request.method == 'POST' and current_user.is_authenticated:
        report = request.get_json()["html_content"]
        report_id = md5(f"{report}{datetime.now()}".encode()).hexdigest()
        user_id = current_user.id
        date = datetime.now()
        title = request.get_json()["html_content_title"]
        
        try:
            new_report = Report(report=report, report_id=report_id, user_id=user_id, date=date, title=title)
            db.session.add(new_report)
            db.session.commit()

        except IntegrityError:
            db.session.rollback()
            flash('Esse relatório já foi salvo.', 'warning')
        
        return {"response":"report saved"}
    
    else:
        return redirect(url_for('main.home'))

@main.route('/delete_report', methods=["POST"])
def delete_report():
    if request.method == "POST" and current_user.is_authenticated:
        report_id = request.get_json()["report_id"]
        report_user_id = Report.query.filter_by(report_id=report_id).first().user_id
        if report_user_id == current_user.id:
            selected_report = Report.query.get(report_id)

            new_trash = Trash(report=selected_report.report, report_id=md5(f"{selected_report.report}{datetime.now()}".encode()).hexdigest(), user_id=selected_report.user_id, date=datetime.now(), title=selected_report.title)
            db.session.add(new_trash)

            db.session.delete(selected_report)
            db.session.commit()

            return {"response":"report deleted"}
    
    else:
        return redirect(url_for('main.home'))

@main.route('/delete_account', methods=["POST"])
def delete_account():
    if request.method == "POST" and current_user.is_authenticated:
        
        reports = Report.query.filter_by(user_id=current_user.id).all()
        for report in reports:
            db.session.delete(report)
        
        trashes = Trash.query.filter_by(user_id=current_user.id).all()
        for trash in trashes:
            db.session.delete(trash)
        
        user = User.query.get_or_404(current_user.id)
        db.session.delete(user)
            
        db.session.commit()

        logout_user()
        flash('Conta apagada com sucesso.', 'success')

        return redirect(url_for('main.home'))
    else:
        return redirect(url_for('main.home'))

@main.route('/reset_code', methods=["POST"])
def reset_code():
    if current_user.is_authenticated:
        global update_account_verification_code
        global code_is_sent

        update_account_verification_code = "0"
        code_is_sent = False
        return {"response":"done"}
    else:
        return redirect(url_for('main.home'))

@main.route('/update_account', methods=["POST"])
def update_account():
    global update_account_verification_code
    global code_is_sent

    if request.method == "POST" and current_user.is_authenticated:
        
        data = request.get_json()

        username = data["username"]
        pwd_old = data["pwd_old"]
        pwd_new = data["pwd_new"]
        email = data["email"]
        code = data["code"]

        is_updating_user = current_user.username != username
        is_updating_password = pwd_old!=pwd_new
        is_updating_email = email!=current_user.email

        if is_updating_user:
            current_user.username = username
            try:
                db.session.commit()

            except IntegrityError:
                db.session.rollback()
            
        if is_updating_password:
            if check_password_hash(current_user.password, pwd_old):
                current_user.password = generate_password_hash(pwd_new)
                try:
                    db.session.commit()
                except IntegrityError as e:
                    db.session.rollback()
                    flash("Erro ao tentar atualizar a conta, tente novamente.", "error")
                    return {"response": "error"}
            else:
                return {"response":"wrong_pwd"} 
        
        if is_updating_email:
            print(code_is_sent)
            print(update_account_verification_code)
            print(code)
            if code_is_sent:
                if code!=update_account_verification_code:
                    print("wrong code")
                    return {"response":"wrong_code"}
                else:
                    current_user.email = email
                    try:
                        db.session.commit()

                    except IntegrityError:
                        db.session.rollback()
                        flash('Email não atualizado. O e-mail já está sendo utilizado.', 'danger')
                        return {"response":"email_already_used"}
                               
                    update_account_verification_code = '0'
                    code_is_sent = False
                    return {"response":"success"}  
            else:
                update_account_verification_code = ''.join(choices(digits, k=6))
                code_is_sent = True
                send_verification_email(email, update_account_verification_code)
                return {"response":"changing_email"} 

        else:
            pass

        flash('Conta atualizada com sucesso.', 'success')
        return {"response":"success"}
    else:
        return {"response":"access denied or error"}

@main.route('/delete_trash', methods=["POST"])
def delete_trash():
    if request.method == "POST":
        trash_id = request.get_json()["report_id"]
        selected_trash = Trash.query.get(trash_id)
        trash_user_id = selected_trash.user_id
        if current_user.is_authenticated and trash_user_id == current_user.id:
            db.session.delete(selected_trash)
            db.session.commit()
            flash('Report apagado com sucesso.', 'success')
            return {"response":"report saved"}
    
    else:
        return redirect(url_for('main.home'))

@main.route("/trash")
def trash():
    if current_user.is_authenticated:
        user_trash = Trash.query.filter_by(user_id=current_user.id).all()
        return render_template('dashboard/trash.html', user_trash=user_trash)

    else:
        return redirect(url_for('main.home'))

@main.route("/recording")
def recording():
    return render_template("recording.html")

@main.route("/privacy_policy")
def privacy_policy():
    return render_template("privacy_policy.html")

@main.route("/terms")
def terms():
    return render_template("terms.html")

@main.route("/history")
def history():
    return render_template("history.html")

@main.route("/security_guide")
def security_guide():
    return render_template("security_guide.html")
